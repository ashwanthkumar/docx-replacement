import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document


def replace_variables(doc: Document, variables: dict[str, str]) -> None:
    """Replace ${VAR_NAME} placeholders in the document with values from the variables dict."""
    pattern = re.compile(r"\$\{(\w+)\}")

    for paragraph in doc.paragraphs:
        _replace_in_runs(paragraph.runs, pattern, variables)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph.runs, pattern, variables)


def _replace_in_runs(runs, pattern: re.Pattern, variables: dict[str, str]) -> None:
    """Replace variables across runs, handling cases where a placeholder spans multiple runs."""
    # Build full text from all runs
    full_text = "".join(run.text for run in runs)
    if not pattern.search(full_text):
        return

    # Try simple per-run replacement first
    simple_full = full_text
    for run in runs:
        new_text = pattern.sub(lambda m: variables.get(m.group(1), m.group(0)), run.text)
        if new_text != run.text:
            run.text = new_text
            simple_full = "".join(r.text for r in runs)

    # Check if all placeholders are resolved after simple replacement
    if not pattern.search(simple_full):
        return

    # Handle placeholders split across runs: rebuild from scratch
    replaced = pattern.sub(lambda m: variables.get(m.group(1), m.group(0)), full_text)
    if runs:
        runs[0].text = replaced
        for run in runs[1:]:
            run.text = ""


def main():
    parser = argparse.ArgumentParser(description="Replace ${VAR_NAME} placeholders in a .docx file")
    parser.add_argument("docx_file", help="Path to the input .docx file")
    parser.add_argument("variables_json", help="Path to a JSON file with variable mappings")
    parser.add_argument("-o", "--output", help="Output .docx file path (default: <input>_filled.docx)")
    args = parser.parse_args()

    docx_path = Path(args.docx_file)
    if not docx_path.exists():
        print(f"Error: {docx_path} not found", file=sys.stderr)
        sys.exit(1)

    json_path = Path(args.variables_json)
    if not json_path.exists():
        print(f"Error: {json_path} not found", file=sys.stderr)
        sys.exit(1)

    with open(json_path) as f:
        variables = json.load(f)

    output_path = args.output or docx_path.with_stem(docx_path.stem + "_filled")

    doc = Document(str(docx_path))
    replace_variables(doc, variables)
    doc.save(str(output_path))
    print(f"Saved to {output_path}")


if __name__ == "__main__":
    main()
